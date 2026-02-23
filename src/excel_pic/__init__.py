from __future__ import annotations

import json
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable

import regex
import typer
from docx import Document
from PIL import Image
import xlsxwriter
from xlsxwriter.exceptions import FileCreateError

DEFAULT_PREFIX_TEXT = """基于五宫格漫画分镜制作二维动漫，严格六宫格分镜布局，全分镜无删减、顺序固定，首帧为Grid0依次到Grid5，镜头切换自然丝滑无卡顿。高饱和配色，色彩浓郁鲜亮且层次分明，人物线条细腻流畅，场景道具还原。人物动作设计充满张力，加入抽帧效果强化动态节奏感，保证整体动态流畅连贯。
重要限制：
1.视频画面禁止出现任何字幕、文字、标题或水印；
2.视频不要有背景音乐，保持静音或仅保留环境音效；
3.Grid0作为首帧，必须是黑色镜头开场；
4.合理运用中景，远景，近景以及特写。
5.中文配音。
注意：角色可以正常说话对话，但视频画面上不要叠加显示任何字幕文字或标题"""

SHOT_RE = re.compile(r"^(?P<episode>\d+)-(?P<scene>\d+)-(?P<shot>\d+)\s+")

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
MAX_EXCEL_CELL_LEN = 32767
MAX_WIN_PATH_LEN = 240
MAX_EXCEL_ROW_HEIGHT = 409
ROW_LINE_HEIGHT_PT = 20
ROW_BASE_PADDING_PT = 16
IMAGE_PADDING_PX = 8
RESERVED_WINDOWS_NAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}

LogSink = Callable[[str, str], None]


class AppError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message
        self.log_path: Path | None = None
        self.report_path: Path | None = None


@dataclass
class LogReport:
    errors: list[dict]
    warnings: list[dict]
    infos: list[str]
    sink: LogSink | None

    def __init__(self, sink: LogSink | None = None) -> None:
        self.errors = []
        self.warnings = []
        self.infos = []
        self.sink = sink

    def _emit(self, level: str, message: str) -> None:
        if self.sink:
            self.sink(level, message)
        else:
            typer.echo(f"[{level}] {message}")

    def info(self, message: str) -> None:
        self.infos.append(message)
        self._emit("INFO", message)

    def warning(self, code: str, message: str, detail: str = "") -> None:
        self.warnings.append({"code": code, "message": message, "detail": detail})
        self._emit("WARN", message)

    def error(self, code: str, message: str, detail: str = "") -> None:
        self.errors.append({"code": code, "message": message, "detail": detail})
        self._emit("ERROR", message)


def detect_runtime_dir() -> Path:
    """Return executable directory when frozen, otherwise current working directory."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path.cwd()


def sanitize_windows_name(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", name)
    cleaned = cleaned.rstrip(" .")
    if cleaned.upper() in RESERVED_WINDOWS_NAMES:
        cleaned = f"_{cleaned}_"
    return cleaned


def derive_export_name_from_images_dir(images_dir: Path) -> str:
    output_name = sanitize_windows_name(images_dir.name)
    if not output_name:
        raise AppError("E007", "输入目录名包含系统不允许的字符，请修改目录名后重试。")
    return output_name


def get_output_paths(exe_dir: Path, images_dir: Path) -> tuple[Path, Path, Path, Path, Path]:
    data_dir = exe_dir / "data"
    config_dir = data_dir / "config"
    export_root = data_dir / "export"
    logs_dir = data_dir / "logs"
    export_name = derive_export_name_from_images_dir(images_dir)
    export_dir = export_root / export_name
    return data_dir, config_dir, export_root, logs_dir, export_dir


def validate_writable_dir(path: Path) -> None:
    try:
        path.mkdir(parents=True, exist_ok=True)
        probe = path / ".write_probe"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
    except Exception as exc:
        raise AppError("E005", "程序目录下 data 不可写，请调整权限后重试。") from exc


def check_path_length(path: Path) -> None:
    if len(str(path)) > MAX_WIN_PATH_LEN:
        raise AppError("E008", "当前路径过长，请把素材放到更短的目录后重试。")


def get_prefix_template(config_dir: Path, prefix_override: str | None) -> str:
    config_dir.mkdir(parents=True, exist_ok=True)
    prefix_file = config_dir / "prompt_prefix.txt"

    if prefix_override is not None:
        prefix_text = prefix_override.strip()
        if not prefix_text:
            raise AppError("E009", "通用前缀模板不能为空，请填写后再生成。")
        prefix_file.write_text(prefix_text, encoding="utf-8")
        return prefix_text

    if not prefix_file.exists():
        prefix_file.write_text(DEFAULT_PREFIX_TEXT, encoding="utf-8")

    prefix_text = prefix_file.read_text(encoding="utf-8").strip()
    if not prefix_text:
        raise AppError("E009", "通用前缀模板不能为空，请填写后再生成。")
    return prefix_text


def collect_docx_lines(word_path: Path, report: LogReport) -> list[str]:
    if word_path.suffix.lower() != ".docx":
        raise AppError("E003", "请选择 .docx 格式的 Word 文件。")
    if not word_path.exists():
        raise AppError("E002", "Word 文件不存在，请重新选择 .docx 文件。")

    doc = Document(word_path)
    lines: list[str] = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            lines.append(txt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    lines.append(txt)

    with zipfile.ZipFile(word_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        if "txbxContent" in xml:
            report.warning("W007", "检测到文本框内容可能未提取，请人工复核。")

    return lines


def parse_episode_id(input_dir_name: str, lines: Iterable[str]) -> str:
    m = re.search(r"(\d+)", input_dir_name)
    if m:
        return m.group(1)

    for line in lines:
        mm = SHOT_RE.match(line)
        if mm:
            return mm.group("episode")

    raise AppError("E004", "没有识别到有效分镜，请检查Word内容格式。")


def extract_scene_groups(lines: Iterable[str], episode_id: str, report: LogReport) -> dict[int, list[str]]:
    grouped: dict[int, list[tuple[int, str]]] = {}
    cross_episode_count = 0

    for line in lines:
        m = SHOT_RE.match(line)
        if not m:
            continue

        ep = m.group("episode")
        if ep != episode_id:
            cross_episode_count += 1
            continue

        scene = int(m.group("scene"))
        shot = int(m.group("shot"))
        grouped.setdefault(scene, []).append((shot, line))

    if cross_episode_count:
        report.warning("W004", "已过滤其他集内容。", detail=f"filtered={cross_episode_count}")

    if not grouped:
        raise AppError("E004", "没有识别到有效分镜，请检查Word内容格式。")

    result: dict[int, list[str]] = {}
    for scene, items in grouped.items():
        sorted_items = sorted(items, key=lambda x: x[0])
        shots = [x[0] for x in sorted_items]
        if shots != list(range(min(shots), max(shots) + 1)):
            report.warning("W003", "分镜编号不连续，已按可识别内容继续生成。", detail=f"scene={scene}")
        result[scene] = [x[1] for x in sorted_items]

    return result


def is_image_readable(path: Path) -> bool:
    try:
        with Image.open(path) as img:
            img.verify()
        return True
    except Exception:
        return False


def copy_images(
    source_dir: Path,
    export_dir: Path,
    strict: bool,
    report: LogReport,
) -> list[Path]:
    if not source_dir.exists() or not source_dir.is_dir():
        raise AppError("E001", "图片源目录不存在，请重新选择。")

    src_images = [p for p in sorted(source_dir.iterdir()) if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
    if not src_images:
        report.warning("W001", "没有找到图片文件，已仅生成文本内容。")
        return []

    copied: list[Path] = []
    name_counter: dict[str, int] = {}

    for p in src_images:
        if not is_image_readable(p):
            msg = "检测到损坏图片，已跳过并记录到报告。"
            if strict:
                raise AppError("E011", msg)
            report.warning("W005", msg, detail=str(p))
            continue

        stem = p.stem
        suffix = p.suffix.lower()
        count = name_counter.get(stem, 0)
        name_counter[stem] = count + 1

        out_name = f"{stem}{suffix}" if count == 0 else f"{stem}_{count+1}{suffix}"
        if count > 0:
            report.warning("W006", "复制图片时出现同名，已自动重命名并记录映射。", detail=f"{p.name} -> {out_name}")

        target = export_dir / out_name
        shutil.copy2(p, target)
        copied.append(target)

    return copied


def map_images_to_scenes(
    images: list[Path],
    episode_id: str,
    scene_numbers: list[int],
    report: LogReport,
) -> dict[int, Path]:
    scene_to_image: dict[int, Path] = {}
    used: set[Path] = set()

    # Primary mapping: parse scene number from filename.
    for img in images:
        stem = img.stem
        match = regex.search(r"(?:(?P<ep>\d+)-)?(?P<scene>\d+)(?!.*\d)", stem)
        if not match:
            continue

        ep = match.group("ep")
        if ep and ep != episode_id:
            continue

        scene = int(match.group("scene"))
        if scene in scene_numbers and scene not in scene_to_image:
            scene_to_image[scene] = img
            used.add(img)

    # Fallback mapping: assign remaining images by scene order.
    remaining_images = [p for p in images if p not in used]
    for scene in scene_numbers:
        if scene in scene_to_image:
            continue
        if remaining_images:
            picked = remaining_images.pop(0)
            scene_to_image[scene] = picked
            used.add(picked)

    missing = [s for s in scene_numbers if s not in scene_to_image]
    extras = [p for p in images if p not in used]

    if missing:
        report.warning("W001", "部分分景没有匹配到图片，已记录到报告。", detail=f"missing_scenes={missing}")
    if extras:
        report.warning("W002", "部分图片未被使用，已记录到报告。", detail=f"unused_images={[p.name for p in extras]}")

    return scene_to_image


def excel_col_width_to_pixels(width: float | None) -> int:
    """Approximate Excel conversion from column-width units to pixels."""
    if width is None:
        width = 8.43
    if width <= 0:
        return 0
    if width < 1:
        return int(width * 12 + 0.5)
    return int(width * 7 + 5)


def pixels_to_points(px: int) -> float:
    return px * 0.75


def estimate_row_text_height_pt(prompt: str) -> float:
    line_count = prompt.count("\n") + 1
    return ROW_BASE_PADDING_PT + line_count * ROW_LINE_HEIGHT_PT


def calc_image_size_by_column(img_path: Path, col_width: float | None) -> tuple[int, int]:
    col_px = excel_col_width_to_pixels(col_width)
    target_w = max(60, col_px - IMAGE_PADDING_PX * 2)
    with Image.open(img_path) as raw:
        src_w, src_h = raw.size

    if src_w <= 0 or src_h <= 0:
        return target_w, max(40, int(target_w * 9 / 16))

    ratio = src_h / src_w
    return target_w, max(40, int(target_w * ratio))


def build_excel(
    excel_path: Path,
    episode_id: str,
    scene_groups: dict[int, list[str]],
    scene_to_image: dict[int, Path],
    prefix_text: str,
) -> None:
    # 先做一轮预计算，避免中途异常时生成半成品 Excel。
    prepared_rows: list[tuple[str, str, Path | None, float]] = []
    c_col_width = 52.89
    for scene in sorted(scene_groups.keys()):
        scene_key = f"{episode_id}-{scene}"
        scene_text = "\n".join(scene_groups[scene])
        prompt = f"{prefix_text}\n{scene_text}"

        if len(prompt) > MAX_EXCEL_CELL_LEN:
            raise AppError("E010", "提示词内容过长，已超过Excel单元格上限，请缩短模板或分镜文本。")

        row_height_pt = estimate_row_text_height_pt(prompt)
        img_path: Path | None = scene_to_image.get(scene)
        if img_path is not None:
            _img_w, img_h = calc_image_size_by_column(img_path, c_col_width)
            row_height_pt = max(row_height_pt, pixels_to_points(img_h + IMAGE_PADDING_PX * 2))

        prepared_rows.append((scene_key, prompt, img_path, min(row_height_pt, MAX_EXCEL_ROW_HEIGHT)))

    wb = xlsxwriter.Workbook(str(excel_path))
    ws = wb.add_worksheet("Sheet1")

    header_fmt = wb.add_format(
        {
            "font_name": "微软雅黑",
            "font_size": 12,
            "bold": True,
            "align": "center",
            "valign": "vcenter",
            "border": 1,
        }
    )
    title_fmt = wb.add_format(
        {
            "font_name": "微软雅黑",
            "font_size": 11,
            "bold": True,
            "align": "center",
            "valign": "vcenter",
            "border": 1,
        }
    )
    prompt_fmt = wb.add_format(
        {
            "font_name": "宋体",
            "font_size": 10,
            "valign": "top",
            "text_wrap": True,
            "border": 1,
        }
    )
    image_cell_fmt = wb.add_format({"border": 1, "valign": "vcenter", "align": "center"})

    ws.set_column("A:A", 11.88)
    ws.set_column("B:B", 88.38)
    ws.set_column("C:C", c_col_width)

    ws.write(0, 0, "视频标题", header_fmt)
    ws.write(0, 1, "提示词", header_fmt)
    ws.write(0, 2, "图片", header_fmt)

    for row_idx, (scene_key, prompt, img_path, row_height_pt) in enumerate(prepared_rows, start=1):
        ws.set_row(row_idx, row_height_pt)
        ws.write(row_idx, 0, scene_key, title_fmt)
        ws.write(row_idx, 1, prompt, prompt_fmt)
        ws.write_blank(row_idx, 2, None, image_cell_fmt)
        if img_path is not None:
            # 使用 Excel 365 的 Place in Cell 语义，避免浮动图层。
            ws.embed_image(row_idx, 2, str(img_path))

    try:
        wb.close()
    except (PermissionError, FileCreateError) as exc:
        raise AppError("E006", "Excel 写入失败，可能被占用，请关闭后重试。") from exc


def write_log_and_report(
    logs_dir: Path,
    input_name: str,
    report: LogReport,
    export_dir: Path | None,
    excel_path: Path | None,
) -> tuple[Path, Path]:
    logs_dir.mkdir(parents=True, exist_ok=True)
    now = datetime.now().strftime("%Y-%m-%d_%H%M%S")

    log_path = logs_dir / f"{now}_{input_name}.log"
    report_path = logs_dir / f"{now}_{input_name}_report.json"

    lines: list[str] = []
    for m in report.infos:
        lines.append(f"[INFO] {m}")
    for w in report.warnings:
        lines.append(f"[WARN] {w['code']} {w['message']} {w.get('detail','')}")
    for e in report.errors:
        lines.append(f"[ERROR] {e['code']} {e['message']} {e.get('detail','')}")

    log_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    report_obj = {
        "status": "failed" if report.errors else "success",
        "export_dir": str(export_dir) if export_dir else None,
        "excel": str(excel_path) if excel_path else None,
        "errors": report.errors,
        "warnings": report.warnings,
    }
    report_path.write_text(json.dumps(report_obj, ensure_ascii=False, indent=2), encoding="utf-8")

    return log_path, report_path


def run_generation(
    images_dir: Path,
    word_file: Path,
    strict: bool,
    prefix_text_override: str | None,
    exe_dir: Path,
    confirm_overwrite: Callable[[Path], bool] | None = None,
    log_sink: LogSink | None = None,
) -> tuple[Path, Path, Path]:
    report = LogReport(sink=log_sink)

    data_dir, config_dir, _export_root, logs_dir, export_dir = get_output_paths(exe_dir, images_dir)
    output_name = export_dir.name or "unknown"
    excel_path: Path | None = None

    try:
        validate_writable_dir(data_dir)
        check_path_length(export_dir)

        prefix_text = get_prefix_template(config_dir, prefix_text_override)

        if export_dir.exists():
            if confirm_overwrite is None:
                answer = typer.confirm(f"导出目录已存在：{export_dir}。是否覆盖？", default=False)
            else:
                answer = confirm_overwrite(export_dir)
            if not answer:
                raise AppError("E012", "已取消覆盖，流程终止。")
            shutil.rmtree(export_dir)

        export_dir.mkdir(parents=True, exist_ok=True)

        report.info(f"读取图片目录：{images_dir}")
        copied_images = copy_images(images_dir, export_dir, strict, report)

        report.info(f"读取Word文件：{word_file}")
        lines = collect_docx_lines(word_file, report)

        episode_id = parse_episode_id(images_dir.name, lines)
        scene_groups = extract_scene_groups(lines, episode_id, report)

        scene_numbers = sorted(scene_groups.keys())
        scene_to_image = map_images_to_scenes(copied_images, episode_id, scene_numbers, report)

        excel_name = f"{episode_id}集.xlsx"
        excel_path = export_dir / excel_name
        check_path_length(excel_path)

        report.info(f"生成Excel：{excel_path}")
        build_excel(excel_path, episode_id, scene_groups, scene_to_image, prefix_text)

        log_path, report_path = write_log_and_report(logs_dir, output_name, report, export_dir, excel_path)
        return export_dir, excel_path, report_path
    except AppError as exc:
        report.error(exc.code, exc.message)
        try:
            log_path, report_path = write_log_and_report(logs_dir, output_name, report, export_dir, excel_path)
            exc.log_path = log_path
            exc.report_path = report_path
        except Exception:
            pass
        raise


def cli(
    images_dir: Path = typer.Option(..., "--images-dir", "-i", file_okay=False, dir_okay=True, resolve_path=True, help="图片源目录"),
    word_file: Path = typer.Option(..., "--word", "-w", file_okay=True, dir_okay=False, resolve_path=True, help="Word路径(.docx)"),
    strict: bool = typer.Option(False, "--strict", help="严格模式：坏图等问题直接中断"),
    prefix_text: str | None = typer.Option(None, "--prefix-text", help="可选：直接传入前缀模板文本"),
    exe_dir: Path = typer.Option(detect_runtime_dir(), "--exe-dir", help="exe同级目录（开发调试可改）"),
) -> None:
    """Generate an Excel file with embedded images from one episode folder."""
    try:
        export_dir, excel_path, report_path = run_generation(
            images_dir=images_dir,
            word_file=word_file,
            strict=strict,
            prefix_text_override=prefix_text,
            exe_dir=exe_dir,
        )
        typer.echo(f"\n完成：\n- 导出目录: {export_dir}\n- Excel: {excel_path}\n- 报告: {report_path}")
    except AppError as exc:
        typer.echo(f"失败 [{exc.code}] {exc.message}")
        if exc.log_path:
            typer.echo(f"- 日志: {exc.log_path}")
        if exc.report_path:
            typer.echo(f"- 报告: {exc.report_path}")
        raise typer.Exit(code=1)


def run() -> None:
    typer.run(cli)


if __name__ == "__main__":
    run()
