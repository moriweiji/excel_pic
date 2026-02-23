from __future__ import annotations

import json
from xml.etree import ElementTree as ET
from zipfile import ZipFile
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from excel_pic import AppError, derive_export_name_from_images_dir, excel_col_width_to_pixels, run_generation


def _create_image(path: Path, color: tuple[int, int, int] = (255, 0, 0)) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (64, 64), color)
    img.save(path)


def _create_docx(path: Path, lines: list[str], table_lines: list[str] | None = None) -> None:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    if table_lines:
        table = doc.add_table(rows=len(table_lines), cols=1)
        for idx, line in enumerate(table_lines):
            table.cell(idx, 0).text = line
    doc.save(path)


def test_run_generation_success(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()
    _create_image(images_dir / "21-1.png", (255, 0, 0))
    _create_image(images_dir / "21-2.png", (0, 255, 0))

    word_file = tmp_path / "剧情分镜.docx"
    _create_docx(
        word_file,
        [
            "第 21 集：测试",
            "21-1-0 [黑屏] 黑色转场，",
            "21-1-1 [全景] 内容1",
            "21-2-0 [黑屏] 黑色转场，",
            "21-2-1 [全景] 内容2",
            "22-1-0 [黑屏] 其他集",
        ],
    )

    exe_dir = tmp_path / "runtime"
    export_dir, excel_path, report_path = run_generation(
        images_dir=images_dir,
        word_file=word_file,
        strict=False,
        prefix_text_override="PREFIX",
        exe_dir=exe_dir,
        confirm_overwrite=lambda _p: True,
    )

    assert export_dir == exe_dir / "data" / "export" / "第21集"
    assert excel_path.exists()
    assert report_path.exists()
    assert not (export_dir / word_file.name).exists()

    wb = load_workbook(excel_path)
    ws = wb.active
    assert ws.cell(1, 1).value == "视频标题"
    assert ws.cell(1, 2).value == "提示词"
    assert ws.cell(1, 3).value == "图片"
    assert ws.cell(2, 1).value == "21-1"
    assert str(ws.cell(2, 2).value).startswith("PREFIX\n")
    assert len(getattr(ws, "_images", [])) == 2

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["status"] == "success"
    assert any(x["code"] == "W004" for x in report["warnings"])


def test_missing_word_writes_failure_report(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()
    _create_image(images_dir / "21-1.png")

    exe_dir = tmp_path / "runtime"
    with pytest.raises(AppError) as ex:
        run_generation(
            images_dir=images_dir,
            word_file=tmp_path / "not-exists.docx",
            strict=False,
            prefix_text_override="PREFIX",
            exe_dir=exe_dir,
            confirm_overwrite=lambda _p: True,
        )

    err = ex.value
    assert err.code == "E002"
    assert err.log_path is not None and err.log_path.exists()
    assert err.report_path is not None and err.report_path.exists()

    report = json.loads(err.report_path.read_text(encoding="utf-8"))
    assert report["status"] == "failed"
    assert any(x["code"] == "E002" for x in report["errors"])


def test_export_name_reserved_word_is_sanitized(tmp_path: Path) -> None:
    images_dir = tmp_path / "CON"
    images_dir.mkdir()
    assert derive_export_name_from_images_dir(images_dir) == "_CON_"


def test_duplicate_image_name_warn(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()
    _create_image(images_dir / "same.png")
    _create_image(images_dir / "same.jpg")
    _create_image(images_dir / "21-1.png")

    word_file = tmp_path / "剧情分镜.docx"
    _create_docx(word_file, ["21-1-0 [黑屏] 黑色转场，", "21-1-1 [全景] 内容1"])

    exe_dir = tmp_path / "runtime"
    _export_dir, _excel_path, report_path = run_generation(
        images_dir=images_dir,
        word_file=word_file,
        strict=False,
        prefix_text_override="PREFIX",
        exe_dir=exe_dir,
        confirm_overwrite=lambda _p: True,
    )

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert any(x["code"] == "W006" for x in report["warnings"])


def test_extract_from_docx_table(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()
    _create_image(images_dir / "21-3.png")

    word_file = tmp_path / "剧情分镜.docx"
    _create_docx(word_file, ["第 21 集"], table_lines=["21-3-0 [黑屏] 黑色转场，", "21-3-1 [全景] 表格内容"])

    exe_dir = tmp_path / "runtime"
    _export_dir, excel_path, _report_path = run_generation(
        images_dir=images_dir,
        word_file=word_file,
        strict=False,
        prefix_text_override="PREFIX",
        exe_dir=exe_dir,
        confirm_overwrite=lambda _p: True,
    )

    wb = load_workbook(excel_path)
    ws = wb.active
    assert ws.cell(2, 1).value == "21-3"
    assert "表格内容" in str(ws.cell(2, 2).value)


def test_too_long_prompt_fails(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()
    _create_image(images_dir / "21-1.png")

    word_file = tmp_path / "剧情分镜.docx"
    _create_docx(word_file, ["21-1-0 [黑屏] 黑色转场，", "21-1-1 [全景] 内容1"])

    exe_dir = tmp_path / "runtime"
    with pytest.raises(AppError) as ex:
        run_generation(
            images_dir=images_dir,
            word_file=word_file,
            strict=False,
            prefix_text_override=("X" * 33000),
            exe_dir=exe_dir,
            confirm_overwrite=lambda _p: True,
        )

    assert ex.value.code == "E010"
    assert ex.value.report_path is not None and ex.value.report_path.exists()


def test_corrupted_image_warns_in_non_strict_mode(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()

    _create_image(images_dir / "21-1.png")
    bad_img = images_dir / "21-2.png"
    bad_img.write_bytes(b"not-an-image")

    word_file = tmp_path / "剧情分镜.docx"
    _create_docx(word_file, ["21-1-0 [黑屏] 黑色转场，", "21-1-1 [全景] 内容1"])

    exe_dir = tmp_path / "runtime"
    _export_dir, _excel_path, report_path = run_generation(
        images_dir=images_dir,
        word_file=word_file,
        strict=False,
        prefix_text_override="PREFIX",
        exe_dir=exe_dir,
        confirm_overwrite=lambda _p: True,
    )

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert any(x["code"] == "W005" for x in report["warnings"])


def test_images_anchor_in_c_column_and_fit_width(tmp_path: Path) -> None:
    images_dir = tmp_path / "第21集"
    images_dir.mkdir()
    _create_image(images_dir / "21-1.png")

    word_file = tmp_path / "剧情分镜.docx"
    _create_docx(word_file, ["21-1-0 [黑屏] 黑色转场，", "21-1-1 [全景] 内容1"])

    exe_dir = tmp_path / "runtime"
    _export_dir, excel_path, _report_path = run_generation(
        images_dir=images_dir,
        word_file=word_file,
        strict=False,
        prefix_text_override="PREFIX",
        exe_dir=exe_dir,
        confirm_overwrite=lambda _p: True,
    )

    wb = load_workbook(excel_path)
    ws = wb.active
    c_col_px = excel_col_width_to_pixels(ws.column_dimensions["C"].width)

    with ZipFile(excel_path) as zf:
        drawing = zf.read("xl/drawings/drawing1.xml")
    root = ET.fromstring(drawing)
    ns = {"xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"}

    from_col = root.find(".//xdr:oneCellAnchor/xdr:from/xdr:col", ns)
    assert from_col is not None and from_col.text == "2"  # 0-based: 2 => C列

    ext = root.find(".//xdr:oneCellAnchor/xdr:ext", ns)
    assert ext is not None
    img_width_px = int(ext.attrib["cx"]) // 9525
    assert img_width_px <= c_col_px
